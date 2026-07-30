#!/usr/bin/env python3
"""
Builds one additional domain-parity fixture case ("entity-resolution-001")
deliberately covering the entity-resolution grouping/ambiguity scenarios
Andrew's Phase 6 instruction asked for: variant-name grouping via a shared
first-name bucket, ambiguous short references between two same-first-name
people, a same-surname-different-first-initial pair that must NOT merge, a
title-prefix candidate that (per the real oracle's own group-key mechanics)
does NOT merge with its untitled counterpart, and a singleton candidate that
must not be grouped with anything.

Synthetic, invented data only, same convention as build_structural_fixtures.py.
Read-only with respect to work/pii_docx_redactor. Reuses build_case() from
export_fixtures.py so this case goes through the identical pipeline and
manifest format as every other domain-parity fixture.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

SCRIPT_DIR = Path(__file__).resolve().parent
DOCSCRUB_WEB = SCRIPT_DIR.parent
REPO_ROOT = DOCSCRUB_WEB.parent
PII_APP_ROOT = REPO_ROOT / "work" / "pii_docx_redactor"

sys.path.insert(0, str(SCRIPT_DIR))
sys.path.insert(0, str(PII_APP_ROOT))

from export_fixtures import build_case  # noqa: E402


def build_entity_resolution_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Entity Resolution Case", level=1)

    # Group A: full-name + comma-form (pre-merged by DETECTION, not entity
    # resolution -- see phase-6-findings.md) + a repeated short first-name
    # reference that entity-resolution groups in via the first-name bucket.
    doc.add_paragraph("Maria Alvarez opened the meeting and introduced the agenda.")
    doc.add_paragraph("Alvarez, Maria then reviewed last quarter's numbers.")
    doc.add_paragraph("Maria confirmed the new deadline soon after.")
    doc.add_paragraph("Later, the team asked Maria for the updated schedule.")

    # Group B/C: two same-first-name people, each with a same-initial
    # variant so each forms a real (2-member) group, plus a repeated bare
    # first name that becomes AMBIGUOUS between the two groups.
    doc.add_paragraph("Andrew Goodloe opened the budget discussion.")
    doc.add_paragraph("Andy Goodloe followed up with additional context.")
    doc.add_paragraph("Andrew Jackson responded with an alternative proposal.")
    doc.add_paragraph("Andy Jackson provided supporting data for that proposal.")
    doc.add_paragraph("Andrew agreed to follow up after the call.")
    doc.add_paragraph("Afterward, everyone thanked Andrew for the clarification.")

    # Must-not-merge: same surname, different first initial.
    doc.add_paragraph("Carlos Mendez signed off on the vendor contract.")
    doc.add_paragraph("Elena Mendez filed the compliance report separately.")

    # Title-prefix quirk: "Dr Susan Whitmore" (no period, so the title token
    # is swept into the detected candidate text) forms a DIFFERENT
    # entity-resolution group key than plain "Susan Whitmore" elsewhere,
    # because the group key uses the candidate's FIRST token as the
    # initial -- "dr" here, not "susan". A real, documented oracle quirk,
    # not a TS deviation -- see phase-6-findings.md.
    doc.add_paragraph("Dr Susan Whitmore presented the quarterly research findings.")
    doc.add_paragraph("Susan Whitmore later answered follow-up questions from the board.")

    # Singleton: appears once, no variant anywhere -- must not be grouped.
    doc.add_paragraph("Priya Natarajan will finalize the vendor selection alone.")

    doc.save(path)


def main() -> None:
    fixtures_root = DOCSCRUB_WEB / "fixtures" / "domain-parity"
    case_id = "entity-resolution-001"
    case_dir = fixtures_root / case_id
    source_dir = case_dir / "source"
    source_dir.mkdir(parents=True, exist_ok=True)
    source_path = source_dir / "synthetic_entity_resolution.docx"
    build_entity_resolution_source(source_path)
    build_case(case_id, source_path, case_dir)


if __name__ == "__main__":
    main()
