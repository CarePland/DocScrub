#!/usr/bin/env python3
"""
Builds three additional domain-parity fixture cases, each deliberately
reproducing one of the structural fidelity risks the OOXML spike found in a
real document (docs/ooxml-spike/findings.md): run-splitting, field codes, and
drawing objects. Closes item 3 of the Phase 1 completion gate in
docs/architecture/phase-1-acceptance-criteria.md.

Uses synthetic, invented data only -- no real documents or real PII are used
or required, consistent with the browser-local / keep-data-local philosophy
the target architecture is built around (architecture v0.2 §4.1). Structural
realism (run fragmentation, field codes, embedded images) is reproduced
directly; the content is fictional.

Read-only with respect to work/pii_docx_redactor: imports its `redactor`
package by path, does not modify it. Reuses build_case() from
export_fixtures.py so all four domain-parity cases go through the identical
pipeline and manifest format.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

SCRIPT_DIR = Path(__file__).resolve().parent
DOCSCRUB_WEB = SCRIPT_DIR.parent
REPO_ROOT = DOCSCRUB_WEB.parent
PII_APP_ROOT = REPO_ROOT / "work" / "pii_docx_redactor"

sys.path.insert(0, str(SCRIPT_DIR))
sys.path.insert(0, str(PII_APP_ROOT))

from export_fixtures import build_case  # noqa: E402


# ---------------------------------------------------------------------------
# Case A: run-splitting -- a candidate name and email fragmented across many
# runs, mirroring the up-to-15-runs-per-paragraph pattern findings.md found
# in a real document (rather than python-docx's normal one-run-per-add_run,
# each fragment here is a handful of characters, matching spell-check-style
# fragmentation).
# ---------------------------------------------------------------------------

def _add_fragmented_run_text(paragraph, text: str, fragment_size: int = 3) -> None:
    for i in range(0, len(text), fragment_size):
        paragraph.add_run(text[i : i + fragment_size])


def build_run_split_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Run-Splitting Case", level=1)

    p = doc.add_paragraph()
    p.add_run("Participant ")
    _add_fragmented_run_text(p, "Priya Natarajan", fragment_size=3)
    p.add_run(" can be reached at ")
    _add_fragmented_run_text(p, "priya.natarajan@example.edu", fragment_size=4)
    p.add_run(" regarding CIN ")
    _add_fragmented_run_text(p, "445566778", fragment_size=2)
    p.add_run(".")

    doc.add_paragraph("A second, unfragmented mention of Priya Natarajan appears here for contrast.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case B: field codes -- a complex field (MERGEFIELD-style) whose *result*
# run contains a candidate name, plus a simple DATE field with no PII, both
# alongside ordinary paragraph text with a second, plain-text candidate.
# ---------------------------------------------------------------------------

def _add_complex_field(paragraph, instruction: str, result_text: str) -> None:
    def _run_with_child(tag: str, **attrs):
        r = OxmlElement("w:r")
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        r.append(el)
        paragraph._p.append(r)
        return r

    _run_with_child("w:fldChar", **{"w:fldCharType": "begin"})

    instr_r = OxmlElement("w:r")
    instr_t = OxmlElement("w:instrText")
    instr_t.set(qn("xml:space"), "preserve")
    instr_t.text = f" {instruction} "
    instr_r.append(instr_t)
    paragraph._p.append(instr_r)

    _run_with_child("w:fldChar", **{"w:fldCharType": "separate"})

    result_r = OxmlElement("w:r")
    result_t = OxmlElement("w:t")
    result_t.text = result_text
    result_r.append(result_t)
    paragraph._p.append(result_r)

    _run_with_child("w:fldChar", **{"w:fldCharType": "end"})


def build_field_codes_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Field Codes Case", level=1)

    p1 = doc.add_paragraph("Prepared for: ")
    _add_complex_field(p1, "MERGEFIELD ParticipantName", "Marcus Whitfield")

    p2 = doc.add_paragraph("Generated on: ")
    _add_complex_field(p2, r'DATE \@ "MM/dd/yyyy"', "01/12/2026")

    doc.add_paragraph(
        "Marcus Whitfield also appears here in plain text, at marcus.whitfield@example.edu, for contrast "
        "with the field-code occurrence above."
    )
    doc.save(path)


# ---------------------------------------------------------------------------
# Case C: drawing objects -- an embedded inline image near candidate text,
# mirroring the pasted-screenshot pattern findings.md flagged as present in
# volume (530 drawing objects) in a real Teams-transcript-style export.
# ---------------------------------------------------------------------------

def _make_placeholder_png() -> bytes:
    from PIL import Image

    img = Image.new("RGB", (64, 64), color=(200, 200, 200))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def build_drawing_objects_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Drawing Objects Case", level=1)

    doc.add_paragraph("Screenshot shared by Diane Okafor (diane.okafor@example.edu) is shown below:")

    png_bytes = _make_placeholder_png()
    image_stream = io.BytesIO(png_bytes)
    doc.add_picture(image_stream)

    doc.add_paragraph("Diane Okafor confirmed the attachment in a follow-up message.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case D: hyperlinks -- a candidate whose display text is inside a
# <w:hyperlink>, plus the same candidate again in plain text for contrast.
# python-docx (this version) has no add_hyperlink() on Paragraph, so this is
# built directly via the low-level oxml API, the same pattern used for field
# codes above.
# ---------------------------------------------------------------------------

def build_hyperlink_source(path: Path) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Hyperlink Case", level=1)

    p = doc.add_paragraph("Contact: ")
    r_id = doc.part.relate_to(
        "mailto:elena.vasquez@example.edu", RT.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = "Elena Vasquez"
    run.append(t)
    hyperlink.append(run)
    p._p.append(hyperlink)

    doc.add_paragraph(
        "Elena Vasquez also appears here in plain text, at elena.vasquez@example.edu, for contrast."
    )
    doc.save(path)


# ---------------------------------------------------------------------------
# Case E: nested tables -- a candidate inside a table nested inside another
# table's cell (depth 2), matching what the structural spike found in a real
# document (max nested table depth 2) and what docx_reader.py already
# explicitly supports reading (iter_paragraphs_in_table recurses into
# cell.tables).
# ---------------------------------------------------------------------------

def build_nested_table_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Nested Table Case", level=1)

    outer = doc.add_table(rows=1, cols=2)
    outer.cell(0, 0).text = "Outer label"
    outer_target_cell = outer.cell(0, 1)
    outer_target_cell.text = ""  # will hold the nested table instead of text

    inner = outer_target_cell.add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "Participant"
    inner.cell(0, 1).text = "Contact"
    inner.cell(1, 0).text = "Grace Kim"
    inner.cell(1, 1).text = "grace.kim@example.edu"

    doc.add_paragraph("Grace Kim also appears here in plain text, outside any table, for contrast.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case F: footer-specific -- a candidate living only in the footer, with no
# occurrence anywhere else, so footer redaction is exercised on its own
# rather than incidentally alongside a header candidate.
# ---------------------------------------------------------------------------

def build_footer_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    section = doc.sections[0]
    section.footer.paragraphs[0].text = "Prepared by Marcus Webb -- confidential"

    doc.add_heading("Synthetic Footer Case", level=1)
    doc.add_paragraph("This document's body contains no candidates at all; the footer does.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case G: diacritics -- NOT a structural risk like the others. Found by
# accident while building the footer case above (an earlier draft used
# "Tomas Reyes" with an accent and the Python oracle detected zero
# candidates for it). Kept as its own fixture specifically to document a
# real, pre-existing gap in the Python reference detector: the fallback
# person-name regex ([A-Z][a-z]{1,30}) is ASCII-only, so any name with a
# non-ASCII letter fails to match as a person candidate at all. This is a
# detection-layer product finding, not an OOXML/rebuild fidelity risk --
# recorded here because "evidence-first engineering" means recording what
# was found, not just what was being looked for.
# ---------------------------------------------------------------------------

def build_diacritics_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Diacritics Case", level=1)
    doc.add_paragraph("Tomás Reyes and José García attended the session.")
    doc.add_paragraph("For contrast, Marcus Webb (ASCII-only) attended as well.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case H: comments -- a candidate in the body (anchored by a comment) plus a
# *different* candidate that exists only inside the comment's own text,
# which lives in a separate OOXML part (word/comments.xml).
# ---------------------------------------------------------------------------

def build_comments_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Comments Case", level=1)

    p = doc.add_paragraph("Reviewed by ")
    run = p.add_run("Harold Ibsen")
    p.add_run(" before submission.")
    doc.add_comment(
        run,
        text="Confirmed with Priscilla Nakamura over email.",
        author="QA Reviewer",
    )

    doc.save(path)


# ---------------------------------------------------------------------------
# Case I: tracked changes -- one candidate inserted (w:ins, ordinary w:t
# inside the wrapper) and one candidate deleted-but-still-present in the
# file (w:del / w:delText -- not visible in Word's normal reading view, but
# physically still in the XML).
# ---------------------------------------------------------------------------

def _add_tracked_insertion(paragraph, text: str, author: str = "Editor") -> None:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-07-27T00:00:00Z")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    ins.append(run)
    paragraph._p.append(ins)


def _add_tracked_deletion(paragraph, text: str, author: str = "Editor") -> None:
    delete = OxmlElement("w:del")
    delete.set(qn("w:id"), "2")
    delete.set(qn("w:author"), author)
    delete.set(qn("w:date"), "2026-07-27T00:00:00Z")
    run = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    del_text.set(qn("xml:space"), "preserve")
    del_text.text = text
    run.append(del_text)
    delete.append(run)
    paragraph._p.append(delete)


def build_tracked_changes_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Tracked Changes Case", level=1)

    p1 = doc.add_paragraph("Newly added reviewer note: ")
    _add_tracked_insertion(p1, "Farrukh Islamov")

    p2 = doc.add_paragraph("Removed from the final version: ")
    _add_tracked_deletion(p2, "Beatrice Alcantara")

    doc.add_paragraph("Farrukh Islamov also appears here in plain text, for contrast.")
    doc.save(path)


# ---------------------------------------------------------------------------
# Case J: text box -- a candidate inside a VML text box (w:pict / v:shape /
# v:textbox / w:txbxContent), the legacy-but-still-common text box format.
# ---------------------------------------------------------------------------

def build_text_box_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Text Box Case", level=1)
    doc.add_paragraph("Body text with no candidates. The text box below holds one.")

    p = doc.add_paragraph()
    run = p.add_run()
    pict = OxmlElement("w:pict")

    shape_xml = (
        '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
        'style="width:200pt;height:40pt" filled="f">'
        "<v:textbox>"
        '<w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:p><w:r><w:t>Note from Winona Featherstone</w:t></w:r></w:p>"
        "</w:txbxContent>"
        "</v:textbox>"
        "</v:shape>"
    )
    shape = parse_xml(shape_xml)
    pict.append(shape)
    run._r.append(pict)

    doc.save(path)


# ---------------------------------------------------------------------------
# Case K: content control -- a candidate inside a structured document tag
# (w:sdt / w:sdtContent), used by Word for form fields and templated
# content.
# ---------------------------------------------------------------------------

def build_content_control_source(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    doc.add_heading("Synthetic Content Control Case", level=1)
    doc.add_paragraph("Body text with no candidates. The content control below holds one.")

    sdt_xml = (
        '<w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:sdtPr><w:alias w:val=\"ParticipantName\"/></w:sdtPr>"
        "<w:sdtContent>"
        "<w:p><w:r><w:t>Desmond Okonkwo</w:t></w:r></w:p>"
        "</w:sdtContent>"
        "</w:sdt>"
    )
    sdt = parse_xml(sdt_xml)
    doc.element.body.insert(len(doc.element.body) - 1, sdt)  # before sectPr

    doc.save(path)


def main() -> None:
    fixtures_root = DOCSCRUB_WEB / "fixtures" / "domain-parity"

    cases = [
        ("run-split-name-001", build_run_split_source, "synthetic_run_split.docx"),
        ("field-codes-001", build_field_codes_source, "synthetic_field_codes.docx"),
        ("drawing-objects-001", build_drawing_objects_source, "synthetic_drawing_objects.docx"),
        ("hyperlink-001", build_hyperlink_source, "synthetic_hyperlink.docx"),
        ("nested-table-001", build_nested_table_source, "synthetic_nested_table.docx"),
        ("footer-001", build_footer_source, "synthetic_footer.docx"),
        ("diacritics-001", build_diacritics_source, "synthetic_diacritics.docx"),
        ("comments-001", build_comments_source, "synthetic_comments.docx"),
        ("tracked-changes-001", build_tracked_changes_source, "synthetic_tracked_changes.docx"),
        ("text-box-001", build_text_box_source, "synthetic_text_box.docx"),
        ("content-control-001", build_content_control_source, "synthetic_content_control.docx"),
    ]

    for case_id, builder, filename in cases:
        case_dir = fixtures_root / case_id
        source_dir = case_dir / "source"
        source_dir.mkdir(parents=True, exist_ok=True)
        source_path = source_dir / filename
        builder(source_path)
        build_case(case_id, source_path, case_dir)


if __name__ == "__main__":
    main()
