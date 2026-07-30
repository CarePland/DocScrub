from pathlib import Path

from docx import Document


def build_sample(path: Path) -> None:
    doc = Document()
    doc.core_properties.author = "Synthetic Author"
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header contact: Jane Smith"
    section.footer.paragraphs[0].text = "Page 1"

    doc.add_heading("Synthetic Teams and Email Transcript", level=1)
    doc.add_paragraph("Jane Smith emailed robert.lee@example.edu about CIN 123456789.")
    doc.add_paragraph("Smith, Jane replied by phone at (555) 123-4567 on 01/12/2026.")
    doc.add_paragraph("Robert Lee mentioned account identifier 998877665544.")
    doc.add_paragraph("Jane Smith and Jane Q Smith are different test names.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Participant"
    table.cell(0, 1).text = "Message"
    table.cell(1, 0).text = "Robert Lee"
    table.cell(1, 1).text = "Please call 555-987-6543."

    doc.save(path)


if __name__ == "__main__":
    build_sample(Path("synthetic_transcript.docx"))
    print("Wrote synthetic_transcript.docx")

