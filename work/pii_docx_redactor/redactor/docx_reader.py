from __future__ import annotations

import hashlib
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from lxml import etree


@dataclass
class TextBlock:
    location: str
    text: str
    paragraph: Paragraph | None = None


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_docx(path: Path) -> DocxDocument:
    return Document(str(path))


def iter_paragraphs_in_table(table: Table, prefix: str) -> Iterable[TextBlock]:
    for row_index, row in enumerate(table.rows, start=1):
        for col_index, cell in enumerate(row.cells, start=1):
            cell_prefix = f"{prefix} table r{row_index}c{col_index}"
            yield from iter_cell_blocks(cell, cell_prefix)


def iter_cell_blocks(cell: _Cell, prefix: str) -> Iterable[TextBlock]:
    for paragraph_index, paragraph in enumerate(cell.paragraphs, start=1):
        yield TextBlock(f"{prefix} paragraph {paragraph_index}", paragraph_text(paragraph), paragraph)
    for table_index, table in enumerate(cell.tables, start=1):
        yield from iter_paragraphs_in_table(table, f"{prefix} nested table {table_index}")


def paragraph_text(paragraph: Paragraph) -> str:
    # Paragraph.text can miss some hyperlink display text in older python-docx
    # versions. Reading w:t nodes directly keeps the scan closer to Word's view.
    nodes = paragraph._element.xpath(".//w:t")
    return "".join(node.text or "" for node in nodes)


def iter_docx_text_blocks(document: DocxDocument) -> List[TextBlock]:
    blocks: List[TextBlock] = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        blocks.append(TextBlock(f"body paragraph {index}", paragraph_text(paragraph), paragraph))
    for table_index, table in enumerate(document.tables, start=1):
        blocks.extend(iter_paragraphs_in_table(table, f"body table {table_index}"))

    for section_index, section in enumerate(document.sections, start=1):
        for index, paragraph in enumerate(section.header.paragraphs, start=1):
            blocks.append(
                TextBlock(
                    f"section {section_index} header paragraph {index}",
                    paragraph_text(paragraph),
                    paragraph,
                )
            )
        for table_index, table in enumerate(section.header.tables, start=1):
            blocks.extend(
                iter_paragraphs_in_table(table, f"section {section_index} header table {table_index}")
            )
        for index, paragraph in enumerate(section.footer.paragraphs, start=1):
            blocks.append(
                TextBlock(
                    f"section {section_index} footer paragraph {index}",
                    paragraph_text(paragraph),
                    paragraph,
                )
            )
        for table_index, table in enumerate(section.footer.tables, start=1):
            blocks.extend(
                iter_paragraphs_in_table(table, f"section {section_index} footer table {table_index}")
            )
    return [block for block in blocks if block.text]


def extract_docx_metadata(path: Path) -> dict[str, str]:
    document = load_docx(path)
    props = document.core_properties
    values = {
        "author": props.author,
        "category": props.category,
        "comments": props.comments,
        "content_status": props.content_status,
        "identifier": props.identifier,
        "keywords": props.keywords,
        "language": props.language,
        "last_modified_by": props.last_modified_by,
        "subject": props.subject,
        "title": props.title,
        "version": props.version,
    }
    return {key: str(value) for key, value in values.items() if value}


def extract_accessible_xml_text(path: Path) -> dict[str, str]:
    """Extract text from selected XML parts that python-docx may not expose."""
    wanted = (
        "word/footnotes.xml",
        "word/endnotes.xml",
        "word/comments.xml",
        "word/document.xml",
    )
    result: dict[str, str] = {}
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        for name in wanted:
            if name not in archive.namelist():
                continue
            root = etree.fromstring(archive.read(name))
            text = "".join(node.text or "" for node in root.xpath(".//w:t", namespaces=namespace))
            if text:
                result[name] = text
    return result

