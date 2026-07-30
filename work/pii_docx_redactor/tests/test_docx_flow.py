from pathlib import Path

from docx import Document

from redactor.decisions import build_default_decisions
from redactor.detectors import detect_all_candidates
from redactor.docx_reader import iter_docx_text_blocks, load_docx
from redactor.docx_writer import rescan_for_originals, write_redacted_docx
from redactor.models import Decision


def build_docx(path: Path):
    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header: Jane Smith"
    section.footer.paragraphs[0].text = "Footer phone 555-222-3333"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Jane Smith")
    run.bold = True
    paragraph.add_run(" emailed jane@example.com.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "CIN 123456789"
    doc.save(path)


def test_text_in_tables_headers_and_footers(tmp_path):
    path = tmp_path / "sample.docx"
    build_docx(path)
    candidates = detect_all_candidates(iter_docx_text_blocks(load_docx(path)), use_spacy=False)
    texts = {candidate.text for candidate in candidates}
    assert "Jane Smith" in texts
    assert "jane@example.com" in texts
    assert "555-222-3333" in texts
    assert "123456789" in texts


def test_output_rescan_and_basic_run_formatting_preserved(tmp_path):
    path = tmp_path / "sample.docx"
    output = tmp_path / "sample_redacted.docx"
    decisions_path = tmp_path / "sample_decisions.json"
    build_docx(path)
    candidates = detect_all_candidates(iter_docx_text_blocks(load_docx(path)), use_spacy=False)
    decisions = build_default_decisions(candidates)
    for candidate in candidates:
        if candidate.text == "Jane Smith":
            decisions[candidate.key].decision = Decision.REDACT
            decisions[candidate.key].replacement = "[PERSON 001]"

    write_redacted_docx(path, output, decisions_path, candidates, decisions)
    redacted = load_docx(output)
    assert redacted.paragraphs[0].runs[0].bold is True
    assert "[PERSON 001]" in redacted.paragraphs[0].text
    remaining = rescan_for_originals(redacted, [c for c in candidates if c.text == "Jane Smith"])
    assert remaining == {}


def test_rename_replaces_original_value(tmp_path):
    path = tmp_path / "sample.docx"
    output = tmp_path / "sample_renamed.docx"
    decisions_path = tmp_path / "sample_decisions.json"
    build_docx(path)
    candidates = detect_all_candidates(iter_docx_text_blocks(load_docx(path)), use_spacy=False)
    decisions = build_default_decisions(candidates)
    for candidate in candidates:
        if candidate.text == "Jane Smith":
            decisions[candidate.key].decision = Decision.RENAME
            decisions[candidate.key].replacement = "Jordan Lee"

    write_redacted_docx(path, output, decisions_path, candidates, decisions)
    renamed = load_docx(output)
    assert "Jordan Lee" in renamed.paragraphs[0].text
    remaining = rescan_for_originals(renamed, [c for c in candidates if c.text == "Jane Smith"])
    assert remaining == {}
