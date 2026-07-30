from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/Local_DOCX_PII_Redactor_User_Guide.docx")


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 108, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Local DOCX PII Redactor")
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(11, 37, 69)
    run.bold = True
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Plain-English user guide for reviewing and redacting Word documents on your Mac")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = MUTED


def add_callout(doc: Document, title: str, body: str, fill=LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [Inches(6.5)])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color="D7DEE8", size="4")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run("•  " + item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(f"{index}.  {item}")


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [Inches(1.875), Inches(4.625)])
    for row_idx, (label, detail) in enumerate(rows):
        label_cell = table.cell(row_idx, 0)
        detail_cell = table.cell(row_idx, 1)
        for cell in (label_cell, detail_cell):
            set_cell_borders(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(label_cell, LIGHT_BLUE)
        label_cell.paragraphs[0].add_run(label).bold = True
        detail_cell.paragraphs[0].add_run(detail)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Local DOCX PII Redactor User Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_footer(doc)
    add_title(doc)
    add_callout(
        doc,
        "Privacy promise",
        "The utility runs on your Mac. It does not use a cloud API, does not use an LLM, and does not change your original Word file.",
        LIGHT_BLUE,
    )

    doc.add_heading("What This Tool Is For", level=1)
    doc.add_paragraph(
        "Use this local app when you have a long Word document, such as Teams or email transcripts, and need help finding possible names, email addresses, phone numbers, and identifying numbers before you manually decide what to redact."
    )
    add_bullets(
        doc,
        [
            "It scans the DOCX and groups repeated values so you do not have to read every page first.",
            "It shows context snippets so you can decide whether each value should stay or be redacted.",
            "It creates a new redacted DOCX plus an audit log and reusable decisions file.",
            "It treats detection as imperfect. You stay in control of every redaction decision.",
        ],
    )

    doc.add_heading("Quick Start", level=1)
    doc.add_paragraph("From Terminal on your Mac:")
    add_label_table(
        doc,
        [
            ("Open folder", "cd /Users/agoodloe/Documents/Codex/2026-07-24/before-making-any-implementation-decisions-read/work/pii_docx_redactor"),
            ("Start app", "python3 -m streamlit run app.py --server.headless true --server.port 8501 --browser.gatherUsageStats false"),
            ("Open browser", "Go to http://localhost:8501"),
        ],
    )
    doc.add_paragraph("The app may already be running. If the page opens, you can start using it immediately.")

    doc.add_heading("Using the App", level=1)
    add_numbered(
        doc,
        [
            "Upload a DOCX file. Keep the original file somewhere safe; the app creates a separate redacted copy.",
            "Click Scan document. The scan runs locally and may take a little while for a long document.",
            "Review each tab: People, Numbers/IDs, Email addresses, Phone numbers, and Other possible identifiers.",
            "For each candidate, choose Keep everywhere, Redact everywhere, Review each occurrence, or leave it Undecided.",
            "Use the context snippets to understand where the value appears without reading the whole document.",
            "If you choose Review each occurrence, decide Keep or Redact for each listed occurrence.",
            "When the safety check looks right, generate the output files.",
        ],
    )

    doc.add_heading("Decision Choices", level=1)
    add_label_table(
        doc,
        [
            ("Keep everywhere", "The value remains unchanged in the redacted output."),
            ("Redact everywhere", "Every matching occurrence is replaced unless an occurrence-specific rule says otherwise."),
            ("Review each occurrence", "You decide occurrence by occurrence. This is useful when the same word is sometimes a name and sometimes not."),
            ("Undecided", "A holding state while you review. The app warns you before output generation."),
        ],
    )

    doc.add_heading("Replacement Labels", level=1)
    doc.add_paragraph(
        "People default to consistent pseudonyms such as [PERSON 001], [PERSON 002], and so on. You can change a replacement label before generating output."
    )
    add_bullets(
        doc,
        [
            "Example: Jane Smith -> [STUDENT 01]",
            "Example: robert.lee@example.edu -> [REDACTED EMAIL]",
            "Example: 123456789 -> [REDACTED ID]",
        ],
    )

    doc.add_heading("Files You Get Back", level=1)
    add_label_table(
        doc,
        [
            ("Redacted DOCX", "<original_name>_redacted.docx"),
            ("Audit log", "<original_name>_redaction_log.csv"),
            ("Reusable decisions", "<original_name>_decisions.json"),
        ],
    )
    doc.add_paragraph(
        "The CSV log records the original candidate, replacement, type, detector source, decision, occurrence count, context, timestamp, and the SHA-256 hash of the input file."
    )

    doc.add_heading("What the Scan Looks At", level=1)
    add_bullets(
        doc,
        [
            "Body paragraphs",
            "Tables, including nested tables where accessible",
            "Headers and footers",
            "Hyperlink display text when Word exposes it as normal text",
            "Core document metadata, listed separately for review",
        ],
    )

    doc.add_heading("Important Limits", level=1)
    add_callout(
        doc,
        "Final review required",
        "Do not treat the output as automatically safe. Open the redacted DOCX in Word and do a final human review before sharing it.",
        "FFF4D6",
    )
    add_bullets(
        doc,
        [
            "Name detection is helpful but imperfect, especially in messy transcripts.",
            "The fallback detector handles First Last and Last, First patterns, but it may miss informal names or unusual capitalization.",
            "Footnotes, endnotes, comments, text boxes, SmartArt, embedded objects, tracked changes, and unusual Word structures may not be redacted perfectly in this MVP.",
            "The tool tries to avoid dates and page numbers when looking for numeric IDs, but you should still review number candidates carefully.",
        ],
    )

    doc.add_heading("Optional Better Name Detection", level=1)
    doc.add_paragraph(
        "If you have Python 3.11+ and want stronger local person-name detection, you can optionally install spaCy and its small English model. This still runs locally."
    )
    add_label_table(
        doc,
        [
            ("Install spaCy", "pip install spacy"),
            ("Install model", "python -m spacy download en_core_web_sm"),
        ],
    )
    doc.add_paragraph(
        "If spaCy is not installed, the app still works using deterministic local patterns."
    )

    doc.add_heading("Testing With Fake Data", level=1)
    doc.add_paragraph("To make a synthetic DOCX with invented names and identifiers:")
    add_label_table(doc, [("Generate sample", "python sample_docx_generator.py")])
    doc.add_paragraph("The sample file is safe test data and does not contain real student information.")

    doc.add_heading("Good Working Habit", level=1)
    add_numbered(
        doc,
        [
            "Scan the original DOCX.",
            "Review and decide candidates.",
            "Generate the redacted DOCX and audit files.",
            "Open the redacted DOCX in Word.",
            "Search the redacted output for any especially sensitive original names or numbers you know about.",
            "Keep the CSV and JSON beside the redacted file for your record.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    main()
